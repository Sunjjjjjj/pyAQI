#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Sep  5 22:11:26 2022

@author: kanonyui
"""
# word_1.py

# 导入库
import sys
# sys.path.insert(0, '/Users/kanonyui/南京浦蓝大气/AQI/')
sys.path.insert(0, '/home/sunji/Scripts/')
sys.path.insert(0, '/home/sunji/Scripts/AQI')
from pyAQI import *
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.ns import qn


# initialization
# outdir = ''
outdir = '/home/sunji/Data/AQI/'
# region
pname = '江苏'
cityname = '盐城市'
citylist = get_cityInProv(pname)
# time period
date = datetime.date.today() - datetime.timedelta(11)
datestr = '%04i年%i月%i日' % (date.year, date.month, date.day)


dateStart = '%04i-%02i-%02i' % (date.year - 1, 1, 1)
dateEnd = str(date)
# crawling data
data = pd.DataFrame()
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    filename = aqi_crawling_city_period(icity['cityName'], dateStart, dateEnd, outdir)
    temp = pd.read_csv(outdir + filename)
    temp['cityName'] = icity['cityName']
    temp['cityId'] = icity['cityId']
    
    data = pd.concat([data, temp])
data.reset_index(inplace = True)
#% get aqi class
inputvar = data[['aqiLevel']].copy()
inputvar.rename(columns = {'aqiLevel': 'level'}, inplace = True)

classes = aqi_class(inputvar)
data = pd.concat([data, classes], axis = 1)
data['datetime'] = pd.to_datetime(data['datetime'], format = '%Y-%m-%d')
#%%
# load data
# data of yesterday
aqi1 = data[data.datetime == '%04i-%02i-%02i' % (date.year, date.month, date.day)]
aqi1.index = aqi1.cityName
# same date from last year 
aqi0 = data[data.datetime == '%04i-%02i-%02i' % (date.year - 1, date.month, date.day)]
aqi0.index = aqi1.cityName
# month
aqimo =  data[(data.datetime >= '%04i-%02i-%02i' % (date.year, date.month, 1))]
# year
aqiyr =  data[data.datetime >= '%04i-%02i-%02i' % (date.year, 1, 1) ]
# data of the same period last year
aqiyr0 =  data[(data.datetime.dt.year == date.year - 1) & (data.datetime <= '%04i-%02i-%02i' % (date.year - 1, date.month, date.day))]


#%% generating report
doc1 = Document()

# set font style
style = doc1.styles['Normal']
style.font.name = 'Times New Roman' # 必须先设置font.name
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 创建段落描述
paragraph = doc1.add_paragraph('各位领导、同事，这是%s空气质量分析简报，请查阅。' % datestr)
doc1.add_paragraph('【截至%s盐城市空气质量情况】' % datestr)
# yesterday aqi
# mask = (aqi1.cityName == cityname)
# idx = mask.index[mask == True]
para1 = doc1.add_paragraph('%s，盐城市环境空气质量为%s，其中PM2.5浓度为%iμg/m3、O3浓度为%iμg/m3,' \
                            '在江苏省13个设区市中分别排名第%i、第%i；'
                            % (datestr, aqi1.loc[cityname]['class'],
                               aqi1.loc[cityname]['pm25'], aqi1.loc[cityname]['o3'],
                               aqi_rank(aqi1, 'pm25').loc[cityname], aqi_rank(aqi1, 'o3').loc[cityname]))
# month aqi
# mask = (aqimo.cityName == cityname)
# idx = mask.index[mask == True]
# qr = aqi_qr(aqimo[mask])

# rank within province
aqimo_m = aqimo.groupby(['cityName']).mean()
aqimo_m['o3'] = aqimo.groupby(['cityName']).quantile(0.9)['o3']
aqimo_m['co'] = aqimo.groupby(['cityName']).quantile(0.95)['co']
aqimo_m['qr'] = np.nan
# qualified rate
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    aqimo_m['qr'].loc[icity.cityName] = aqi_qr(aqimo[aqimo.cityName == icity.cityName])

para1.add_run('%i月1-%i日，盐城市PM2.5浓度均值为%1.1fμg/m3，O3浓度均值为%1.1fμg/m3，优良率为%1.1f%%，'\
              '在江苏省13个设区市中分别排名第%i、第%i、第%i；'
              % (date.month, date.day, aqimo_m.loc[cityname]['pm25'], aqimo_m.loc[cityname]['o3'], 
                 aqimo_m.loc[cityname]['qr'],
                 aqi_rank(aqimo_m, 'pm25').loc[cityname], aqi_rank(aqimo_m, 'o3').loc[cityname], 
                 aqi_rank(aqimo_m, 'qr', ascending = False).loc[cityname]))
# year aqi
# rank within province
aqiyr_m = aqiyr.groupby(['cityName']).mean()
aqiyr_m['o3'] = aqiyr.groupby(['cityName']).quantile(0.9)['o3']
aqiyr_m['co'] = aqiyr.groupby(['cityName']).quantile(0.95)['co']
aqiyr_m['qr'] = np.nan
# qualified rate
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    aqiyr_m['qr'].loc[icity.cityName] = aqi_qr(aqiyr[aqiyr.cityName == icity.cityName])
    
para1.add_run('截至%s，盐城市PM2.5浓度均值为%1.1fμg/m3，O3浓度均值为%1.1fμg/m3，优良率为%1.1f%%，'\
              '在江苏省13个设市中分别排名第%i、第%i、第%i；'
              % (datestr, aqiyr_m.loc[cityname]['pm25'], aqiyr_m.loc[cityname]['o3'],
                 aqiyr_m.loc[cityname]['qr'],
                 aqi_rank(aqiyr_m, 'pm25').loc[cityname], aqi_rank(aqiyr_m, 'o3').loc[cityname],
                 aqi_rank(aqiyr_m, 'qr', ascending = False).loc[cityname]))

# compare with last year
# rank within province
aqiyr0_m = aqiyr0.groupby(['cityName']).mean()
aqiyr0_m['o3'] = aqiyr0.groupby(['cityName']).quantile(0.9)['o3']
aqiyr0_m['co'] = aqiyr0.groupby(['cityName']).quantile(0.95)['co']
aqiyr0_m['qr'] = np.nan
# qualified rate
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    aqiyr0_m['qr'].loc[icity.cityName] = aqi_qr(aqiyr0[aqiyr0.cityName == icity.cityName])

dpm25 = (aqiyr_m['pm25'] - aqiyr0_m['pm25']) / aqiyr0_m['pm25'] * 1e2
do3 = (aqiyr_m['o3'] - aqiyr0_m['o3']) / aqiyr0_m['o3'] * 1e2
dqr = aqiyr_m['qr'] - aqiyr0_m['qr']


para1.add_run('和去年同期相比，分别%s、%s、%s百分点，'\
              '恶化程度在江苏省13个设区市中分别排名第%i、第%i、第%i。'
              % (val2str(dpm25.loc[cityname].round(1)), val2str(do3.loc[cityname].round(1)),
                 val2str(dqr.loc[cityname].round(1)), 
                 dpm25.rank(ascending = False, method = 'min').loc[cityname],
                 do3.rank(ascending = False, method = 'min').loc[cityname],
                 dqr.rank(ascending = False, method = 'min').loc[cityname]))
                   
#%%
target_pm25 = 28
target_qr = 87
days_in_year = len(pd.date_range('%04i-01-01' % (date.year), '%04i-12-31' % (date.year), freq = '1D'))
rest_days = len(pd.date_range(date, '%04i-12-31' % (date.year), freq = '1D'))

# 清洁日向上取整，污染日向下取整
target_pol_days = np.ceil((1 - target_qr / 1e2) * days_in_year)
already_pol_days = np.ceil((1 - aqiyr_m.loc[cityname]['qr'] / 1e2) * (days_in_year - rest_days))
planned_polluted_days = np.floor(max(0, (target_pol_days - already_pol_days)))

date1 = date + pd.to_timedelta(1, unit = 'day')
datestr1 = '%04i年%i月%i日' % (date1.year, date1.month, date1.day)

rest_pm25 = (target_pm25 * days_in_year - aqiyr_m.loc[cityname]['pm25'] * (days_in_year - rest_days)) / rest_days


para2 = doc1.add_paragraph('【盐城市空气质量达标规划】')
para2 = doc1.add_paragraph('%04i年盐城市PM2.5累积均值目标为%1.1fμg/m3。'\
                    '截至%s，PM2.5累积均值为%1.1fμg/m3。'
                    % (date.year, target_pm25, datestr, aqiyr_m.loc[cityname]['pm25']))
para2.add_run('为实现年度目标，%s起剩余%i日全市6个国控站点PM2.5日均浓度应低于%1.1fμg/m3。'\
                    '%04i年盐城市优良天数比率目标为%i%%。截至%s，累积优良天数比率为%1.1f%%。'
                    % (datestr1, rest_days - 1, rest_pm25, 
                       date.year, target_qr, datestr, aqiyr_m.loc[cityname]['qr']))


if aqiyr_m.loc[cityname]['qr'] >= target_qr:
    para2.add_run('为实现年度目标，%s起剩余%i日最少还需%i个优良天，最多允许超标%i天。'
                        % (datestr1, rest_days - 1, rest_days - planned_polluted_days, planned_polluted_days))
else:
    para2.add_run('为实现年度目标，已超标%i天。'
                 % (abs(target_pol_days - already_pol_days)))


#%%
# data from web
columns = ['so2', 'no2', 'pm10', 'co', 'o3_8h', 'pm25']
snames = ['大丰高级中学(国)', '月亮广场', '宝龙广场', '盐城电厂', '市监测站', '开发区管委会']

# data = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/六项浓度均值.xls', encoding = 'gbk', header = 1)[0]
# data.drop(index = 0, inplace = True)
# data = pd.DataFrame(data.iloc[:, [1, 5, 9, 13, 17, 19]].values, index = data.iloc[:, 0], columns = columns)
# data = data.astype(float)
# data = data.loc[snames]

data = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/优良率2022.xls', encoding = 'gbk', header = 1)[0]
data.index = data.iloc[:, 0]
data = data.loc[snames]
data['优良率'] = [float(i[:-1]) for i in data['优良率']]

data0 = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/优良率2021.xls', encoding = 'gbk', header = 1)[0]
data0.index = data0.iloc[:, 0]
data0 = data0.loc[snames]
data0['优良率'] = [float(i[:-1]) for i in data0['优良率']]

columns = ['PM2.5均值', 'O3-8日均值90%位数', '优良率']
diff = (data[columns] - data0[columns]) / data0[columns] * 1e2
diff['优良率'] =  data['优良率'] - data0['优良率']


#%%
var = 'PM2.5均值'
data['pm25排名'] = aqi_rank(data, var, ascending = False, method = 'min')
sites = data.sort_values('pm25排名').index[:3].to_list()

doc1.add_paragraph('【截至%s国控点空气质量情况】' % datestr)
para3 = doc1.add_paragraph('今年，PM2.5倒数3名为：%s、%s、%s，'\
                    '浓度分别为%1.1fμg/m3、%1.1fμg/m3、%1.1fμg/m3；'
                    % (sites[0], sites[1], sites[2],
                       data[var].loc[sites[0]], data[var].loc[sites[1]], data[var].loc[sites[2]]))

var = 'O3-8日均值90%位数'
data['o3排名'] = aqi_rank(data, var, ascending = False, method = 'min')
sites = data.sort_values('o3排名').index[:3].to_list()
para3.add_run('O3倒数为：%s、%s、%s，'\
                '浓度分别为%1.1fμg/m3、%1.1fμg/m3、%1.1fμg/m3；'
                % (sites[0], sites[1], sites[2],
                   data[var].loc[sites[0]], data[var].loc[sites[1]], data[var].loc[sites[2]]))

var = '优良率'
data['优良率排名'] = aqi_rank(data, var, ascending = True, method = 'min')
sites = data.sort_values('优良率排名').index[:3].to_list()
para3.add_run('优良率倒数3名为：%s、%s、%s，'\
                    '分别为%1.1f%%、%1.1f%%、%1.1f%%。'
                    % (sites[0], sites[1], sites[2],
                        data[var].loc[sites[0]], data[var].loc[sites[1]], data[var].loc[sites[2]]))
    
# compare with last year
var = 'PM2.5均值'
diff['pm25排名'] = aqi_rank(diff, var, ascending = False, method = 'min')
num_deg = (diff.sort_values('pm25排名').iloc[:3][var] > 0).sum()
sites = diff.sort_values('pm25排名').index[:min(3, num_deg)].to_list()
para3.add_run('PM2.5同比恶化程度前%i名为：%s、%s、%s，'\
              '分别同比上升%1.1f%%、%1.1f%%、%1.1f%%；'
              % (num_deg, sites[0], sites[1], sites[2],
                diff[var].loc[sites[0]], diff[var].loc[sites[1]], diff[var].loc[sites[2]]))

var = 'O3-8日均值90%位数'
diff['o3排名'] = aqi_rank(diff, var, ascending = False, method = 'min')
num_deg = (diff.sort_values('o3排名').iloc[:3][var] > 0).sum()
sites = diff.sort_values('o3排名').index[:3].to_list()
para3.add_run('O3同比恶化程度前%s名为：%s、%s、%s，'\
              '分别同比上升%1.1f%%、%1.1f%%、%1.1f%%；'
              % (num_deg, sites[0], sites[1], sites[2],
               diff[var].loc[sites[0]], diff[var].loc[sites[1]], diff[var].loc[sites[2]]))
    
var = '优良率'
diff['优良率排名'] = aqi_rank(diff, var, ascending = True, method = 'min')
num_deg = (diff.sort_values('优良率').iloc[:3][var] < 0).sum()
sites = diff.sort_values('优良率排名').index[:3].to_list()
para3.add_run('优良率同比恶化前%i名为：%s、%s、%s，'\
              '分别同比下降%1.1f%%、%1.1f%%、%1.1f%%。'
              % (num_deg, sites[0], sites[1], sites[2],
                 abs(diff[var].loc[sites[0]]), abs(diff[var].loc[sites[1]]), abs(diff[var].loc[sites[2]])))

# 保存文件
doc1.save('日报推送_%04i-%02i-%02i_test.docx' % (date.year, date.month, date.day))
print('完成！')