#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Sep  5 22:11:26 2022

@author: kanonyui
"""
# word_1.py

# 导入库
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.ns import qn

# date of yesterday
date = datetime.date.today() - datetime.timedelta(1)
datestr = '%i月%i日' % (date.month, date.day)

# create new file
doc1 = Document()

# set font style
style = doc1.styles['Normal']
style.font.name = 'Times New Roman' # 必须先设置font.name
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# 创建段落描述
paragraph = doc1.add_paragraph('各位领导、同事，这是%s空气质量分析简报，请查阅。' % datestr)

doc1.add_paragraph('【截至%s盐城市空气质量情况】' % datestr)
doc1.add_paragraph('%s，盐城市环境空气质量优，其中PM2.5浓度为6μg/m3、O3浓度为117μg/m3,' \
                   '在江苏省13个设区市中分别排名第2、第10；'\
                   '8月1-24日，盐城市PM2.5浓度均值为18.5μg/m3，'\
                   'O3浓度均值为183μg/m3，优良率为75.0%%，'\
                   '在江苏省13个设区市中分别排名第4、第8、第6；'\
                   '截至8月23日，盐城市PM2.5浓度均值为28.5μg/m3，'\
                   'O3浓度均值为180μg/m3，优良率为79.7%%，'\
                   '在江苏省13个设市中分别排名第2、第8、第2；'\
                   '和去年同期相比，分别上升10.0%%、上升16.1%%、下降6.3%%，'\
                   '恶化程度在江苏省13个设区市中分别排名第2、第2、第11。'
                    % (datestr))
doc1.add_paragraph('【盐城市空气质量达标规划】')
doc1.add_paragraph('2022年盐城市PM2.5累积均值目标为28 μg/m3。'\
                   '截至%s，PM2.5累积均值为28.5μg/m3。'\
                   '为实现年度目标，8月24日起剩余129日全市6个国控站点PM2.5日均浓度应低于27.1μg/m3。'\
                   '2022年盐城市优良天数比率目标为87%%。截至8月24日，累积优良天数比率为79.7%%。'\
                   '为实现年度目标，8月24日起剩余129日最少还需129个优良天，最多允许超标0天。'
                   % (datestr) )
doc1.add_paragraph('【截至%s国控点空气质量情况】' % datestr )
doc1.add_paragraph('今年，PM2.5倒数3名为：盐城电厂、市监测站、月亮广场，'\
                   '浓度分别为31.5μg/m3、30.6μg/m3、28.9μg/m3；'\
                   'O3倒数为：宝龙广场、盐城电厂、大丰高级中学，'\
                   '浓度分别为185μg/m3、183μg/m3、181μg/m3；'\
                   '优良率倒数3名为：盐城电厂、大丰高级中学、月亮广场，'\
                   '分别为75.8%%、77.0%%、77.4%%。'\
                   'PM2.5同比恶化程度前3名为：月亮广场、市监测站、盐城电厂，'\
                   '分别同比上升20.9%%、17.7%%、10.5%%；'\
                   'O3同比恶化程度前3名为：盐城电厂、宝龙广场、大丰高级中学，'\
                   '分别同比上升21.2%%、18.6%%、17.5%%；'\
                   '优良率同比恶化前3名为：盐城电厂、大丰高级中学、月亮广场，'\
                   '分别同比下降9.8%%、9.2%%、8.3%%。')



# 保存文件
doc1.save('日报推送_%04i-%02i-%02i.docx' % (date.year, date.month, date.day))