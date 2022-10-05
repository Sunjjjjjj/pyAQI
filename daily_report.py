#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Sep 21 14:42:35 2022

@author: sunji
"""

# 导入库
import sys
sys.path.insert(0, '/home/sunji/Scripts/')
sys.path.insert(0, '/home/sunji/Scripts/AQI')
from pyAQI import *
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.ns import qn

# initialization
outdir = '/home/sunji/Data/AQI/'
# region
pname = '江苏'
cityname = '盐城市'
citylist = get_cityInProv(pname)
# time period
date = datetime.date.today() - datetime.timedelta(11)
datestr = '%04i年%i月%i日' % (date.year, date.month, date.day)

# 空气质量发布APP数据爬虫
dateStart = '%04i-%02i-%02i' % (date.year - 1, 1, 1)
dateEnd = str(date)


data = pd.DataFrame()
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    filename = aqi_crawling_city_period(icity['cityName'], dateStart, dateEnd, outdir)
    temp = pd.read_csv(outdir + filename)
    temp['cityName'] = icity['cityName']
    temp['cityId'] = icity['cityId']
    
    data = pd.concat([data, temp])
data.reset_index(inplace = True)
#% get aqi class
inputvar = data[['aqiLevel']].copy()
inputvar.rename(columns = {'aqiLevel': 'level'}, inplace = True)

classes = aqi_class(inputvar)
data = pd.concat([data, classes], axis = 1)
data['datetime'] = pd.to_datetime(data['datetime'], format = '%Y-%m-%d')
#%%
# load data
# data of yesterday
aqi1 = data[data.datetime == '%04i-%02i-%02i' % (date.year, date.month, date.day)]
aqi1.index = aqi1.cityName
# same date from last year 
aqi0 = data[data.datetime == '%04i-%02i-%02i' % (date.year - 1, date.month, date.day)]
aqi0.index = aqi1.cityName
# month
aqimo =  data[(data.datetime >= '%04i-%02i-%02i' % (date.year, date.month, 1))]
# year
aqiyr =  data[data.datetime >= '%04i-%02i-%02i' % (date.year, 1, 1) ]
# data of the same period last year
aqiyr0 =  data[(data.datetime.dt.year == date.year - 1) & (data.datetime <= '%04i-%02i-%02i' % (date.year - 1, date.month, date.day))]

#%% generating report
doc1 = Document()

# set font style
style = doc1.styles['Normal']
style.font.name = 'Times New Roman' # 必须先设置font.name
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 创建段落描述
para1 = doc1.add_paragraph('1.盐城市环境空气质量情况')
para1 = doc1.add_paragraph('%i月%i日，盐城市环境空气质量为%s，其中PM2.5浓度为%iμg/m3、O3浓度为%iμg/m3,' \
                            '在江苏省13个设区市中分别排名第%i、第%i；'
                            % (date.month, date.day, aqi1.loc[cityname]['class'],
                               aqi1.loc[cityname]['pm25'], aqi1.loc[cityname]['o3'],
                               aqi_rank(aqi1, 'pm25').loc[cityname], aqi_rank(aqi1, 'o3').loc[cityname]))
# rank within province
aqimo_m = aqimo.groupby(['cityName']).mean()
aqimo_m['o3'] = aqimo.groupby(['cityName']).quantile(0.9)['o3']
aqimo_m['co'] = aqimo.groupby(['cityName']).quantile(0.95)['co']
aqimo_m['qr'] = np.nan
# qualified rate
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    aqimo_m['qr'].loc[icity.cityName] = aqi_qr(aqimo[aqimo.cityName == icity.cityName])

para1.add_run('%i月1-%i日，盐城市PM2.5浓度均值为%1.1fμg/m3，O3浓度均值为%1.1fμg/m3，优良率为%1.1f%%，'\
              '在江苏省13个设区市中分别排名第%i、第%i、第%i；'
              % (date.month, date.day, aqimo_m.loc[cityname]['pm25'], aqimo_m.loc[cityname]['o3'], 
                 aqimo_m.loc[cityname]['qr'],
                 aqi_rank(aqimo_m, 'pm25').loc[cityname], aqi_rank(aqimo_m, 'o3').loc[cityname], 
                 aqi_rank(aqimo_m, 'qr', ascending = False).loc[cityname]))
# year aqi
# rank within province
aqiyr_m = aqiyr.groupby(['cityName']).mean()
aqiyr_m['o3'] = aqiyr.groupby(['cityName']).quantile(0.9)['o3']
aqiyr_m['co'] = aqiyr.groupby(['cityName']).quantile(0.95)['co']
aqiyr_m['qr'] = np.nan
# qualified rate
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    aqiyr_m['qr'].loc[icity.cityName] = aqi_qr(aqiyr[aqiyr.cityName == icity.cityName])
    
para1.add_run('截至%s，盐城市PM2.5浓度均值为%1.1fμg/m3，O3浓度均值为%1.1fμg/m3，优良率为%1.1f%%，'\
              '在江苏省13个设市中分别排名第%i、第%i、第%i；'
              % (datestr, aqiyr_m.loc[cityname]['pm25'], aqiyr_m.loc[cityname]['o3'],
                 aqiyr_m.loc[cityname]['qr'],
                 aqi_rank(aqiyr_m, 'pm25').loc[cityname], aqi_rank(aqiyr_m, 'o3').loc[cityname],
                 aqi_rank(aqiyr_m, 'qr', ascending = False).loc[cityname]))

# compare with last year
# rank within province
aqiyr0_m = aqiyr0.groupby(['cityName']).mean()
aqiyr0_m['o3'] = aqiyr0.groupby(['cityName']).quantile(0.9)['o3']
aqiyr0_m['co'] = aqiyr0.groupby(['cityName']).quantile(0.95)['co']
aqiyr0_m['qr'] = np.nan
# qualified rate
for i in range(len(citylist)):
    icity = citylist.iloc[i]
    aqiyr0_m['qr'].loc[icity.cityName] = aqi_qr(aqiyr0[aqiyr0.cityName == icity.cityName])

dpm25 = (aqiyr_m['pm25'] - aqiyr0_m['pm25']) / aqiyr0_m['pm25'] * 1e2
do3 = (aqiyr_m['o3'] - aqiyr0_m['o3']) / aqiyr0_m['o3'] * 1e2
dqr = aqiyr_m['qr'] - aqiyr0_m['qr']


para1.add_run('和去年同期相比，分别%s、%s、%s百分点，'\
              '恶化程度在江苏省13个设区市中分别排名第%i、第%i、第%i。'
              % (val2str(dpm25.loc[cityname].round(1)), val2str(do3.loc[cityname].round(1)),
                 val2str(dqr.loc[cityname].round(1)), 
                 dpm25.rank(ascending = False, method = 'min').loc[cityname],
                 do3.rank(ascending = False, method = 'min').loc[cityname],
                 dqr.rank(ascending = False, method = 'min').loc[cityname]))
#%% annual data
table1 = pd.DataFrame()
table2 = pd.DataFrame()
table3 = pd.DataFrame()
table4 = pd.DataFrame()


columns = ['so2', 'no2', 'pm10', 'co', 'o3_8h', 'pm25', 'AQI']
csites = ['盐城电厂', '市监测站', '宝龙广场', '开发区管委会', '大丰高级中学(国)', '月亮广场']

data_d = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/AQI统计(2022-09-19).xls', encoding = 'gbk', header = 4)[0]
data_d.drop(index = [0, len(data_d) - 1], inplace = True)
data_d = pd.DataFrame(data_d.iloc[:, [1, 3, 5, 7, 9, 11, 13]].values, index = data_d.iloc[:, 0], columns = columns)
data_d = data_d.astype(float)

data_ds = data_d.loc[csites]
for ivar in columns:
    data_ds[ivar+'排名'] = aqi_rank(data_ds, ivar, method = 'min')



data_mo = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/优良率202209.xls', encoding = 'gbk', header = 1)[0]
data_mo.index = data_mo.iloc[:, 0]
data_mo['优良率'] = [float(i[:-1]) for i in data_mo['优良率']]

data_mos = data_mo.loc[csites]
for ivar in ['SO2均值', 'NO2均值', 'PM10均值', 'CO日均值95%位数', 'O3-8日均值90%位数', 'PM2.5均值']:
    data_mos[ivar+'排名'] = aqi_rank(data_mos, ivar, method = 'min')
data_mos['优良率排名'] = aqi_rank(data_mos, '优良率', method = 'min', ascending = False)



data = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/优良率2022.xls', encoding = 'gbk', header = 1)[0]
data.index = data.iloc[:, 0]
data['优良率'] = [float(i[:-1]) for i in data['优良率']]
data_s = data.loc[csites]
for ivar in ['SO2均值', 'NO2均值', 'PM10均值', 'CO日均值95%位数', 'O3-8日均值90%位数', 'PM2.5均值']:
    data_s[ivar+'排名'] = aqi_rank(data_s, ivar, method = 'min')
data_s['优良率排名'] = aqi_rank(data_s, '优良率', method = 'min', ascending = False)


data0 = pd.read_html('/home/sunji/Documents/盐城市大气污染防治/Data/每日推送数据/优良率2021.xls', encoding = 'gbk', header = 1)[0]
data0.index = data0.iloc[:, 0]
data0['优良率'] = [float(i[:-1]) for i in data0['优良率']]
data0_s = data0.loc[csites]
for ivar in ['SO2均值', 'NO2均值', 'PM10均值', 'CO日均值95%位数', 'O3-8日均值90%位数', 'PM2.5均值']:
    data0_s[ivar+'排名'] = aqi_rank(data0_s, ivar, method = 'min')
data0_s['优良率排名'] = aqi_rank(data0_s, '优良率', method = 'min', ascending = False)


columns = ['PM2.5均值', 'O3-8日均值90%位数', '优良率']
diff = (data[columns] - data0[columns]) / data0[columns] * 1e2
diff['优良率'] =  data['优良率'] - data0['优良率']

diff_s = diff.loc[csites]
for ivar in diff.columns:
    diff_s[ivar+'排名'] = aqi_rank(diff_s, ivar, method = 'min', ascending = False)
diff_s['优良率排名'] = aqi_rank(diff_s, '优良率', method = 'min')


table1 = pd.concat([table1, data_ds])
table2 = pd.concat([table2, data_mos])
table3 = pd.concat([table3, data_s])
table4 = pd.concat([table4, diff_s])
#%%
para2 = doc1.add_paragraph('2.6个国控站点环境空气质量情况')

for isite in csites:
    para2 = doc1.add_paragraph('%s，%s国控站点PM2.5、O3浓度分别为%iμg/m3、%iμg/m3，在全市6个国控站点中分别排名第%i、第%i。'\
                               '%i月1-%i日，PM2.5、O3浓度均值及优良率分别为%1.1fμg/m3、%1.1fμg/m3、%1.1f%%，在国控站点中分别排名第%i、第%i、第%i。'\
                               '截至%s，PM2.5、O3浓度及优良率分别为%1.1fμg/m3、%1.1fμg/m3、%1.1f%%，在国控站点中分别排名第%i、第%i、第%i；'\
                               '和去年同期相比，分别%s、%s、%s，恶化程度分别排名第%i、第%i、第%i。' 
                               % (datestr[5:], isite, data_ds.loc[isite]['pm25'], data_ds.loc[isite]['o3_8h'],
                                  data_ds.loc[isite]['pm25排名'], data_ds.loc[isite]['o3_8h排名'], 
                                  date.month, date.day, data_mos.loc[isite]['PM2.5均值'], data_mos.loc[isite]['O3-8日均值90%位数'], data_mos.loc[isite]['优良率'], data_mos.loc[isite]['PM2.5均值排名'], data_mos.loc[isite]['O3-8日均值90%位数排名'], data_mos.loc[isite]['优良率排名'],
                                  datestr[5:],data_s.loc[isite]['PM2.5均值'], data_s.loc[isite]['O3-8日均值90%位数'], data_s.loc[isite]['优良率'], data_s.loc[isite]['PM2.5均值排名'], data_s.loc[isite]['O3-8日均值90%位数排名'], data_s.loc[isite]['优良率排名'],
                                  
                                  val2str(diff_s.loc[isite]['PM2.5均值'].round(1)), val2str(diff_s.loc[isite]['O3-8日均值90%位数'].round(1)), val2str(diff_s.loc[isite]['优良率'].round(1)), diff_s.loc[isite]['PM2.5均值排名'], diff_s.loc[isite]['O3-8日均值90%位数排名'], diff_s.loc[isite]['优良率排名']
                                  ))
# 保存文件

#%%
psites = ['盐都盐塘河公园', '盐都区郭猛社区', '大丰创意产业园','响水职业中学', '响水县自来水厂','滨海司法局', '滨海中专','阜宁中学','阜宁金沙湖街道', '射阳环保局','射阳县监测站','建湖县书画院', '建湖二中','东台西溪植物园','东台实中南校区']

data_ds = data_d.loc[psites]
for ivar in data_ds.columns:
    data_ds[ivar+'排名'] = aqi_rank(data_ds, ivar, method = 'min')
    
data_mos = data_mo.loc[psites]
for ivar in ['SO2均值', 'NO2均值', 'PM10均值', 'CO日均值95%位数', 'O3-8日均值90%位数', 'PM2.5均值']:
    data_mos[ivar+'排名'] = aqi_rank(data_mos, ivar, method = 'min')
data_mos['优良率排名'] = aqi_rank(data_mos, '优良率', method = 'min', ascending = False)

data_s = data.loc[psites]
for ivar in ['SO2均值', 'NO2均值', 'PM10均值', 'CO日均值95%位数', 'O3-8日均值90%位数', 'PM2.5均值']:
    data_s[ivar+'排名'] = aqi_rank(data_s, ivar, method = 'min')
data_s['优良率排名'] = aqi_rank(data_s, '优良率', method = 'min', ascending = False)

data0_s = data.loc[psites]
for ivar in ['SO2均值', 'NO2均值', 'PM10均值', 'CO日均值95%位数', 'O3-8日均值90%位数', 'PM2.5均值']:
    data0_s[ivar+'排名'] = aqi_rank(data0_s, ivar, method = 'min')
data0_s['优良率排名'] = aqi_rank(data0_s, '优良率', method = 'min', ascending = False)

diff_s = diff.loc[psites]
for ivar in diff.columns:
    diff_s[ivar+'排名'] = aqi_rank(diff_s, ivar, method = 'min', ascending = False)
diff_s['优良率排名'] = aqi_rank(diff_s, '优良率', method = 'min')



table1 = pd.concat([table1, data_ds])
table2 = pd.concat([table2, data_mos])
table3 = pd.concat([table3, data_s])
table4 = pd.concat([table4, diff_s])

def func(data, var, rank, pct = False):
    temp = data.sort_values('%s排名' % var)[data['%s排名' % var] <= rank]
    sites = ''
    values = ''
    ranks = ''
    
    if pct == True:
        unit = '%'
    else:
        unit = 'μg/m3'
    
    for isite in temp.index:
        sites = sites + isite + '、'
        values = values + '%1.1f%s、' % (abs(temp.loc[isite][var]), unit)
        ranks = ranks + '第%i、' % temp.loc[isite]['%s排名' % var]
    
    return sites, values, ranks

#%%

para3 = doc1.add_paragraph('3.15个省控站点环境空气质量情况')

sites, values, ranks = func(data_ds, 'pm25', 3)
para3 = doc1.add_paragraph('%i月%i日，全市15个省控站点中，%sPM2.5浓度分别为%s，分别排名%s；' % (date.month, date.day, sites[:-1], values[:-1], ranks[:-1]))


    
sites, values, ranks = func(data_ds, 'o3_8h', 3)
para3.add_run('%sO3浓度分别为%s，分别排名%s。'
              % (sites[:-1], values[:-1], ranks[:-1]))


sites, values, ranks = func(data_mos, 'PM2.5均值', 3)
para3 = doc1.add_paragraph('%i月1-%i日，%sPM2.5浓度分别为%s，分别排名%s；'
                           % (date.month, date.day, sites[:-1], values[:-1], ranks[:-1]))


sites, values, ranks = func(data_mos, 'O3-8日均值90%位数', 3)
para3.add_run('%sO3浓度分别为%s，分别排名%s。'
              % (sites[:-1], values[:-1], ranks[:-1]))


    
sites, values, ranks = func(data_s, 'PM2.5均值', 3)
para3 = doc1.add_paragraph('截至%i月%i日，%sPM2.5浓度分别为%s，分别排名%s；'
                           %(date.month, date.day, sites[:-1], values[:-1], ranks[:-1]))

sites, values, ranks = func(data_s, 'O3-8日均值90%位数', 3)
para3.add_run('%sO3浓度分别为%s，分别排名%s；'
              % (sites[:-1], values[:-1], ranks[:-1]))

sites, values, ranks = func(data_s, '优良率', 3, pct = True)
para3.add_run('%s优良率分别为%s，分别排名%s。' 
              % (sites[:-1], values[:-1], ranks[:-1]))


#%%


sites, values, ranks = func(diff_s[diff_s['PM2.5均值'] > 0], 'PM2.5均值', 3, pct = True)
if len(sites) > 0:
    para3.add_run('%sPM2.5同比分别上升%s，恶化程度分别为%s；'
                  % (sites[:-1], values[:-1], ranks[:-1]))


sites, values, ranks = func(diff_s[diff_s['O3-8日均值90%位数'] > 0], 'O3-8日均值90%位数', 3, pct = True)
if len(sites) > 0:
    para3.add_run('%sO3同比分别上升%s，恶化程度分别为%s；'
                  % (sites[:-1], values[:-1], ranks[:-1]))
else:
    para3.add_run('')


sites, values, ranks = func(diff_s[diff_s['优良率'] < 0], '优良率', 3, pct = True)
if len(sites) > 0:
    para3.add_run('%s优良率同比分别下降%s，恶化程度分别为%s。'
                  % (sites[:-1], values[:-1], ranks[:-1]))
else:
    para3.add_run('')


#%%
target_pm25 = 28
target_qr = 87
days_in_year = len(pd.date_range('%04i-01-01' % (date.year), '%04i-12-31' % (date.year), freq = '1D'))
rest_days = len(pd.date_range(date, '%04i-12-31' % (date.year), freq = '1D'))

# 清洁日向上取整，污染日向下取整
target_pol_days = np.ceil((1 - target_qr / 1e2) * days_in_year)
already_pol_days = np.ceil((1 - aqiyr_m.loc[cityname]['qr'] / 1e2) * (days_in_year - rest_days))
planned_polluted_days = np.floor(max(0, (target_pol_days - already_pol_days)))

date1 = date + pd.to_timedelta(1, unit = 'day')
datestr1 = '%04i年%i月%i日' % (date1.year, date1.month, date1.day)


rest_pm25 = (target_pm25 * days_in_year - aqiyr_m.loc[cityname]['pm25'] * (days_in_year - rest_days)) / rest_days

para4 = doc1.add_paragraph('5.可达性分析')
para4 = doc1.add_paragraph('PM2.5：%i年盐城市PM2.5累积均值目标为%1.1fμg/m3。截至%i月%i日，PM2.5累积均值为%1.1fμg/m3。为实现年度目标，%s起剩余%i日全市6个国控站点PM2.5日均浓度应低于%1.1fμg/m3。'
                           % (date.year, target_pm25, date.month, date.day, aqiyr_m.loc[cityname]['pm25'],
                              datestr1, rest_days-1, rest_pm25))
para4.add_run('优良天数比率：%i年盐城市优良天数比率目标为%i。截至%s，累积优良天数比率为%1.1f%%。为实现年度目标，%s起剩余%i日最少还需%i个优良天，最多允许超标%i天。'
              % (date.year, target_qr, datestr, aqiyr_m.loc[cityname]['qr'], datestr1, rest_days - 1, rest_days - 1 - planned_polluted_days, planned_polluted_days ))

doc1.save('日报_%04i-%02i-%02i_test.docx' % (date.year, date.month, date.day))

#%% tables
table1 = table1[['pm25', 'pm25排名', 'o3_8h', 'o3_8h排名', 'AQI', 'AQI排名']].astype(int)
table1.columns = pd.MultiIndex.from_tuples([('PM2.5', r'浓度(μg/m3)'), ('PM2.5', '排名'),
                           ('O3_8h', r'浓度(μg/m3)'), ('O3_8h', '排名'), 
                           ('AQI', r'(-)'), ('AQI', '排名')])


table2['优良率'] = table2['优良率'] / 1e2
table2 = table2[['PM2.5均值', 'PM2.5均值排名', 'O3-8日均值90%位数', 'O3-8日均值90%位数排名', '优良率', '优良率排名']]
# table2.iloc[:, [0,1,2,3, 5]] = table2.iloc[:, [0,1,2,3, 5]].astype(int)
table2.columns = pd.MultiIndex.from_tuples([('PM2.5', r'浓度(μg/m3)'), ('PM2.5', '排名'),
                            ('O3_8h', r'浓度(μg/m3)'), ('O3_8h', '排名'), 
                            ('优良率', r'(%)'), ('优良率', '排名')])


table3['优良率'] = table3['优良率'] / 1e2
table4.iloc[:, :3] = table4.iloc[:, :3] / 1e2
for ikey in table4.columns:
    table3['%s同比' % ikey] = table4[ikey]
    # table4.rename(columns = {ikey: '%s-1' % (ikey, inplace = True})
table3 = table3.drop(columns = '名称')

table3 = table3[['PM2.5均值', 'PM2.5均值排名', 'PM2.5均值同比', 'PM2.5均值排名同比', 
                 'O3-8日均值90%位数', 'O3-8日均值90%位数排名', 'O3-8日均值90%位数同比', 'O3-8日均值90%位数排名同比', 
                 '优良率', '优良率排名', '优良率同比', '优良率排名同比']]
# table2.iloc[:, [0,1,2,3, 5]] = table2.iloc[:, [0,1,2,3, 5]].astype(int)
table3.columns = pd.MultiIndex.from_tuples([('PM2.5', r'浓度(μg/m3)'), ('PM2.5', '排名'), ('PM2.5', r'同比情况'), ('PM2.5', '恶化排名'),
                            ('O3_8h', r'浓度(μg/m3)'), ('O3_8h', '排名'), ('O3_8h', r'同比情况'), ('O3_8h', '恶化排名'),
                            ('优良率', r'(%)'), ('优良率', '排名'), ('优良率', r'同比情况'), ('优良率', '恶化排名')])

                            
#%%
writer = pd.ExcelWriter('%i年%i月%i日日报表格.xlsx' % (date.year, date.month, date.day))
table1.to_excel(writer, sheet_name = '表-1 %s盐城市个站点环境空气质量表' % (datestr), index_label = '监测站点', merge_cells = True)
table2.to_excel(writer, sheet_name = '表-2 %i年%i月1日-%i日盐城市个站点环境空气质量表' % (date.year, date.month, date.day), index_label = '监测站点', merge_cells = True)
table3.to_excel(writer, sheet_name = '表-3 截至%s盐城市个站点环境空气质量表' % (datestr), index_label = '监测站点', merge_cells = True)

workbook = writer.book
format1 = workbook.add_format({'num_format': '0.0%'})
format_red = workbook.add_format({'bg_color': '#FFC7CE'})
format_green = workbook.add_format({'bg_color': '#E2EFDA'})
format_yellow = workbook.add_format({'bg_color': '#FFF2CC'})
format_noborder = workbook.add_format({'bottom': 0, 'top': 0, 'left': 0, 'right': 0})
format_border = workbook.add_format({'bottom': 1, 'top': 1, 'left': 1, 'right': 1})

worksheet = writer.sheets['表-1 %s盐城市个站点环境空气质量表' % (datestr)]
worksheet.set_column(0, 0, 12)
worksheet.set_column(1, table1.shape[1], 10)
worksheet.conditional_format(0, 0, table1.shape[0] + 2, table1.shape[1],
                             {'type': 'no_errors',
                                 'format': format_border})

worksheet.conditional_format(3, 0, 8, 0,
                             {'type': 'cell',
                              'criteria': '>',
                              'value': 3,
                              'format': format_yellow})

for icol in [2, 4, 6]:
    worksheet.conditional_format(3, icol, 8, icol,
                                 {'type': 'cell',
                                  'criteria': '>',
                                  'value': 3,
                                  'format': format_red})
    worksheet.conditional_format(9, icol, table1.shape[0] + 2, icol,
                                 {'type': 'cell',
                                  'criteria': '<',
                                  'value': 4,
                                  'format': format_green})
    worksheet.conditional_format(9, icol, table1.shape[0] + 2, icol,
                                 {'type': 'cell',
                                  'criteria': '>',
                                  'value': 12,
                                  'format': format_red})

worksheet = writer.sheets['表-2 %i年%i月1日-%i日盐城市个站点环境空气质量表' % (date.year, date.month, date.day)]
worksheet.set_column(5, 5, 10, format1)
worksheet.set_column(0, 0, 12)
worksheet.set_column(1, table2.shape[1], 10)
worksheet.conditional_format(0, 0, table2.shape[0] + 2, table2.shape[1],
                             {'type': 'no_errors',
                                 'format': format_border})

worksheet.conditional_format(3, 0, 8, 0,
                             {'type': 'cell',
                              'criteria': '>',
                              'value': 3,
                              'format': format_yellow})
for icol in [2, 4, 6]:
    worksheet.conditional_format(3, icol, 8, icol,
                                 {'type': 'cell',
                                  'criteria': '>',
                                  'value': 3,
                                  'format': format_red})
    worksheet.conditional_format(9, icol, table1.shape[0] + 2, icol,
                                 {'type': 'cell',
                                  'criteria': '<',
                                  'value': 4,
                                  'format': format_green})
    worksheet.conditional_format(9, icol, table1.shape[0] + 2, icol,
                                 {'type': 'cell',
                                  'criteria': '>',
                                  'value': 12,
                                  'format': format_red})



worksheet = writer.sheets['表-3 截至%s盐城市个站点环境空气质量表' % (datestr)]
worksheet.set_column(3, 3, 10, format1)
worksheet.set_column(7, 7, 10, format1)
worksheet.set_column(11, 11, 10, format1)

worksheet.set_column(0, 0, 12)
worksheet.set_column(1, table3.shape[1], 10)
worksheet.conditional_format(0, 0, table3.shape[0] + 2, table3.shape[1],
                             {'type': 'no_errors',
                                 'format': format_border})

worksheet.conditional_format(3, 0, 8, 0,
                             {'type': 'cell',
                              'criteria': '>',
                              'value': 3,
                              'format': format_yellow})
for icol in np.arange(2, 12, 2):
    worksheet.conditional_format(3, icol, 8, icol,
                                 {'type': 'cell',
                                  'criteria': '>',
                                  'value': 3,
                                  'format': format_red})
    worksheet.conditional_format(9, icol, table1.shape[0] + 2, icol,
                                 {'type': 'cell',
                                  'criteria': '<',
                                  'value': 4,
                                  'format': format_green})
    worksheet.conditional_format(9, icol, table1.shape[0] + 2, icol,
                                 {'type': 'cell',
                                  'criteria': '>',
                                  'value': 12,
                                  'format': format_red})

worksheet.conditional_format(3, 12, table1.shape[0] + 2, 12,
                             {'type': 'cell',
                              'criteria': '<',
                              'value': 4,
                              'format': format_red})

writer.save()
print('完成！')
