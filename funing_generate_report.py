#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Sep 16 14:02:23 2022

@author: sunji
"""

import datetime
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def aqi_rank(inputdata, var, method = 'dense', ascending = True):
    return inputdata[var].rank(ascending = ascending, method =  method).astype(int)
    

def aqi_qr(inputdata):
    return ((inputdata['class'] == '优') | (inputdata['class'] == '良')).sum() / len(inputdata) * 1e2
    
def val2str(value, pct = True):
    if value < 0:
        valstr = '下降' + str(abs(value))
    if value == 0:
        valstr = '无变化'
    if value > 0:
        valstr = '升高' + str(value)
    
    if (value != 0) & (pct == True):
        valstr = valstr + '%'
    return valstr

def generate_report():
    # initialization
    outdir = '/home/sunji/Documents/阜宁月报/'
    # region
    pname = '江苏'
    cityname = '盐城市'
    dname = '阜宁县'
    # time period
    date = datetime.date.today() - datetime.timedelta(30)
    datestr = '%04i年%02i月%02i日' % (date.year, date.month, date.day)
    print('制作%04i年%02i月报告...' % (date.year, date.month))
    
    # 六项浓度均值(全站点全区县)
    data_tb = pd.read_html('Data/六项浓度均值tb.xls', encoding = 'gbk', header = 1)[0]
    data_hb = pd.read_html('Data/六项浓度均值hb.xls', encoding = 'gbk', header = 1)[0]
    data = pd.read_html('Data/六项浓度均值.xls', encoding = 'gbk', header = 1)[0]
    
    # 去年-今年每月类别统计
    data_cls = pd.read_html('Data/月类别统计.xls', encoding = 'gbk', header = 0)[0]
    # 阜宁每日数据
    ddata = pd.read_html('Data/AQI统计(阜宁县).xls', encoding = 'gbk', header = 4)[0]
    
    # pre-processing
    columns = ['so2', 'no2', 'pm10', 'co', 'o3_8h', 'pm25']
    # 六项浓度均值(全站点全区县)
    data_tb.drop_duplicates(subset = '名称', keep = 'last', inplace = True)
    data_hb.drop_duplicates(subset = '名称', keep = 'last', inplace = True)
    data.drop_duplicates(subset = '名称', keep = 'last', inplace = True)
    
    data_tb.drop(index = 0, inplace = True)
    data_hb.drop(index = 0, inplace = True)
    data.drop(index = 0, inplace = True)
    
    data_tb = pd.DataFrame(data_tb.iloc[:, [1, 5, 9, 13, 17, 19]].values, index = data_tb.iloc[:, 0], columns = columns)
    data_hb = pd.DataFrame(data_hb.iloc[:, [1, 5, 9, 13, 17, 19]].values, index = data_hb.iloc[:, 0], columns = columns)
    data = pd.DataFrame(data.iloc[:, [1, 5, 9, 13, 17, 19]].values, index = data.iloc[:, 0], columns = columns)
    
    data_tb = data_tb.astype(float)
    data_hb = data_hb.astype(float)
    data = data.astype(float)
    
    # 去年-今年每月类别统计
    data_cls.drop(index = 0, inplace = True)
    data_cls.drop(data_cls[data_cls['月份'] == '汇总'].index, inplace = True)
    data_cls.index = data_cls['月份']
    data_cls.iloc[:, 2:13] = data_cls.iloc[:, 2:13].astype(float)
    
    # 阜宁每日数据
    ddata = pd.DataFrame(ddata.iloc[:, [0, 1, 3, 5, 7, 9, 11, 14, 15]].values, columns = ['日期'] + columns + ['首要污染物', '级别'])
    ddata.index = ddata['日期']
    #%%
    # create new file
    doc1 = Document()
    
    # set font style
    style = doc1.styles['Normal']
    style.font.name = 'Times New Roman' # 必须先设置font.name
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    
    hb = round((data - data_hb) / data_hb * 1e2, 2)
    tb = round((data - data_tb) / data_tb * 1e2, 2)
    
    
    data_tmp = data_cls[data_cls['站点'] == dname].copy()
    yl = data_tmp.loc['%04i年%02i月' % (date.year, date.month)]['优'] + data_tmp.loc['%04i年%02i月' % (date.year, date.month)]['良']
    ylrate = data_tmp.loc['%04i年%02i月' % (date.year, date.month)]['达标率']
    yl0 = data_tmp.loc['%04i年%02i月' % (date.year, date.month -1)]['优'] + data_tmp.loc['%04i年%02i月' % (date.year, date.month - 1)]['良']
    ylrate0 = data_tmp.loc['%04i年%02i月' % (date.year, date.month -1)]['达标率']
    
    
    para1 = doc1.add_paragraph('一、8月空气质量概况')
    para1 = doc1.add_paragraph('%i月阜宁县PM2.5平均浓度%1.1f μg/m³，上月平均浓度%1.1f μg/m³，环比%s，同比%s；'\
                       'PM10本月平均浓度%i μg/m³，上月平均浓度%i μg/m³，环比%s，同比%s；'\
                       'O3本月平均浓度%i μg/m³，上月平均浓度%i μg/m³，环比%s，同比%s；'\
                       '本月优良天%i天，优良率%s。上月优良天%i天，优良率%s。'
                       % (date.month, data.loc[dname, 'pm25'], data_hb.loc[dname, 'pm25'], val2str(hb.loc[dname, 'pm25']), val2str(tb.loc[dname, 'pm25']), 
                          data.loc[dname, 'pm10'], data_hb.loc[dname, 'pm25'], val2str(hb.loc[dname, 'pm10']), val2str(tb.loc[dname, 'pm10']),
                          data.loc[dname, 'o3_8h'], data_hb.loc[dname, 'o3_8h'], val2str(hb.loc[dname, 'o3_8h']), val2str(tb.loc[dname, 'o3_8h']),
                          yl, ylrate, yl0, ylrate0))
    
    #%% table 1 
    dnames = ['阜宁县','建湖县', '开发区', '盐都区', '城南新区', '亭湖区', '东台市', '射阳县', '响水县', '大丰县', '滨海县']
    
    table1 = pd.DataFrame(data.loc[dnames][['pm10', 'pm25', 'o3_8h']].values, columns = ['pm10', 'pm25', 'O3_8h'], index = dnames)
    table1['pm10排名'] = aqi_rank(table1, 'pm10', method = 'min').values
    table1['pm25排名'] = aqi_rank(table1, 'pm25', method = 'min').values
    table1['O3_8h排名'] = aqi_rank(table1, 'O3_8h', method = 'min').values
    table1 = table1[sorted(table1.columns)]
    table1.index.name = '区县'
    table1.to_excel('表1-%04i-%02i盐城市区县空气质量排名.xlsx' % (date.year, date.month))
    
    
    #%%
    dsites = ['阜宁中学','阜宁金沙湖街道']
    snames = ['盐都盐塘河公园', '盐都区郭猛社区', '大丰创意产业园','响水职业中学', '响水县自来水厂','滨海司法局', '滨海中专','阜宁中学','阜宁金沙湖街道', '射阳环保局','射阳县监测站','建湖县书画院', '建湖二中','东台西溪植物园','东台实中南校区']
    distrcts = ['盐都高新区', '盐都区', '大丰区', '响水县', '响水县', '滨海县', '滨海县', '阜宁县', '阜宁县', '射阳县', '射阳县', '建湖县', '建湖县', '东台市','东台市']
    
    pdata = data.loc[snames]
    para2 = doc1.add_paragraph('%i月阜宁县省控监测点位情况如下：'\
                              '阜宁中学点位PM2.5、PM10质量浓度在全市15个监测点位分别位列为第%i名、第%i名；'\
                              'O3为第%i名，NO2为第%i名，SO2为第%i名，CO为第%i名。'\
                              '阜宁金沙湖街道点位PM2.5、PM10质量浓度在全市15个监测点位中分别位列第%i名、第%i名；'\
                              'O3为第%i名，NO2为第%i名；SO2为第%i名，CO为第%i名。'
                  % (date.month, 
                     aqi_rank(pdata, 'pm25', method = 'min').loc[dsites[0]], 
                     aqi_rank(pdata, 'pm10', method = 'min').loc[dsites[0]],
                     aqi_rank(pdata, 'o3_8h', method = 'min').loc[dsites[0]],
                     aqi_rank(pdata, 'no2', method = 'min').loc[dsites[0]],
                     aqi_rank(pdata, 'so2', method = 'min').loc[dsites[0]],
                     aqi_rank(pdata, 'co', method = 'min').loc[dsites[0]],
                     aqi_rank(pdata, 'pm25', method = 'min').loc[dsites[1]], 
                     aqi_rank(pdata, 'pm10', method = 'min').loc[dsites[1]],
                     aqi_rank(pdata, 'o3_8h', method = 'min').loc[dsites[1]],
                     aqi_rank(pdata, 'no2', method = 'min').loc[dsites[1]],
                     aqi_rank(pdata, 'so2', method = 'min').loc[dsites[1]],
                     aqi_rank(pdata, 'co', method = 'min').loc[dsites[1]],
                     ))
    #%% table 2
    table2 = pd.DataFrame(pdata.values, columns = ['SO2', 'NO2', 'pm10', 'CO', 'O3_8h', 'pm25'])
    
    
    for ikey in table2.columns:
        table2[ikey + '排名'] = aqi_rank(table2, ikey)
    
    table2 = table2[sorted(table2.columns)]
    table2['站点'] = snames
    table2['区县'] = distrcts
    
    
    table2.index = pd.MultiIndex.from_tuples(list(zip(distrcts, snames)))
    table2.drop(['区县', '站点'], axis = 1, inplace = True)
    table2.to_excel('表2-%04i-%02i盐城市省控点空气质量排名.xlsx' % (date.year, date.month))
    #%%
    tb = round((data - data_tb) / data_tb * 1e2, 2)
    
    data_cls['优良天'] = data_cls['优'] + data_cls['良']
    data_tmp = data_cls.loc['%04i年%02i月' % (date.year, date.month)].copy()
    
    data_tmp[data_tmp['站点'] == dsites[0]]
    para2 = doc1.add_paragraph('(一) %s' % dsites[0])
    para2 = doc1.add_paragraph('本月%s站点出现优良天数%i天（有效天数%i天）；'\
                               '本月PM2.5平均浓度%1.2f μg/m³，同比2021年%s；'\
                               '本月PM10平均浓度%1i μg/m³，同比2021年%s；'\
                               '本月O3平均浓度%i μg/m³，同比2021年%s。'
                        % (dsites[0], data_tmp[data_tmp['站点'] == dsites[0]]['优良天'], data_tmp[data_tmp['站点'] == dsites[0]]['有效天数'],
                           data['pm25'].loc[dsites[0]], val2str(tb['pm25'].loc[dsites[0]]),
                           data['pm10'].loc[dsites[0]], val2str(tb['pm10'].loc[dsites[0]]),
                           data['o3_8h'].loc[dsites[0]], val2str(tb['o3_8h'].loc[dsites[0]]),
                            ))
    para2 = doc1.add_paragraph('(二) %s' % dsites[1])
    para2 = doc1.add_paragraph('本月%s站点出现优良天数%i天（有效天数%i天）；'\
                               '本月PM2.5平均浓度%1.2f μg/m³，同比2021年%s；'\
                               '本月PM10平均浓度%1i μg/m³，同比2021年%s；'\
                               '本月O3平均浓度%i μg/m³，同比2021年%s。'
                        % (dsites[1], data_tmp[data_tmp['站点'] == dsites[1]]['优良天'], data_tmp[data_tmp['站点'] == dsites[1]]['有效天数'],
                           data['pm25'].loc[dsites[1]], val2str(tb['pm25'].loc[dsites[1]]),
                           data['pm10'].loc[dsites[1]], val2str(tb['pm10'].loc[dsites[1]]),
                           data['o3_8h'].loc[dsites[1]], val2str(tb['o3_8h'].loc[dsites[1]]),
                            ))
    
        
    pday = ddata[ddata['级别'] == '三级']
    dates = ''
    prmpol = ''
    for i in pday.index.tolist():
        dates = dates + i + '、'
        prmpol = prmpol + pday.loc[i]['首要污染物'] + '、'
        
        
    
    para2 = doc1.add_paragraph('(三) %i月污染天回顾' % date.month)
    para2 = doc1.add_paragraph('%04i年%02i月，阜宁市空气质量为三级共%i天(%s)，首要污染物为%s。'
                        % (date.year, date.month, len(pday), dates[:-1], prmpol[:-1]
                            ))
    
    #%%
    fig = plt.figure(figsize = (8, 3))
    ax1 = fig.add_axes([0.1, 0.3, 0.8, 0.6])
    plt.plot(ddata['日期'], ddata['pm25'],'k', label = 'pm2.5')
    plt.plot(ddata['日期'], ddata['pm10'], 'k--', label = 'pm10')
    plt.ylabel(r'颗粒物[$\mu g/m^3$]', font = 'SimHei')
    plt.legend(loc = 2, frameon = False, ncol = 2)
    plt.xticks(rotation = 45, fontsize = 8)
    plt.xlim(ddata['日期'].min(), ddata['日期'].max())
    ax2 = plt.twinx()
    plt.plot(ddata['日期'], ddata['o3_8h'], c = 'r', label = 'O3_8h')
    ax2.tick_params(axis ='y', labelcolor = 'r')
    ax2.set_ylabel('臭氧[$\mu g/m^3$]', color = 'r', font = 'SimHei')
    plt.legend(frameon = False)
    
    plt.savefig('图1-%04i-%02i%spm25-pm10-o3_8h时序图.png' % (date.year, date.month, dname))
    
    
    doc1.save('%04i年%02i月阜宁月报.docx' % (date.year, date.month))
    
    print('完成!')